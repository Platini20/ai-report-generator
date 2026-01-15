"""
Module d'export en format Word (.docx) avec graphiques intégrés
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from typing import Dict, Any, Optional
import matplotlib.pyplot as plt


def add_figure_to_doc(doc: Document, fig: plt.Figure, width: float = 6.0):
    """
    Ajoute une figure matplotlib au document Word
    
    Args:
        doc: Document Word
        fig: Figure matplotlib
        width: Largeur en inches
    """
    try:
        # Sauvegarder la figure dans un buffer
        buffer = io.BytesIO()
        fig.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        buffer.seek(0)
        
        # Ajouter au document
        doc.add_picture(buffer, width=Inches(width))
        plt.close(fig)
        
    except Exception as e:
        print(f"Error adding figure to Word doc: {e}")


def generate_word_report(df, analysis: Dict[str, Any], ai_insights: Dict[str, Any], 
                        lang: str = 'fr', visualizations: Optional[Dict] = None) -> bytes:
    """
    Génère un rapport Word (.docx) complet avec graphiques
    
    Args:
        df: DataFrame pandas
        analysis: Dictionnaire d'analyse
        ai_insights: Insights générés
        lang: Langue ('fr' ou 'en')
        visualizations: Dict {nom: (figure, interprétation)} optionnel
        
    Returns:
        bytes: Contenu du fichier Word
    """
    doc = Document()
    
    # Style du document
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Traductions
    titles = {
        'fr': {
            'report': 'RAPPORT D\'ANALYSE PROFESSIONNEL',
            'summary': 'Résumé Exécutif',
            'overview': 'Vue d\'Ensemble',
            'stats': 'Statistiques Détaillées',
            'visualizations': 'Visualisations & Analyses',
            'cat_dist': 'Distribution des Catégories',
            'trends': 'Tendances Principales',
            'analysis': 'Analyses Approfondies',
            'reco': 'Recommandations Stratégiques',
            'conclusion': 'Conclusion',
            'obs': 'Observations',
            'vars': 'Variables',
            'num': 'Numériques',
            'cat': 'Catégorielles',
            'interpretation': 'Interprétation',
        },
        'en': {
            'report': 'PROFESSIONAL ANALYSIS REPORT',
            'summary': 'Executive Summary',
            'overview': 'Overview',
            'stats': 'Detailed Statistics',
            'visualizations': 'Visualizations & Analysis',
            'cat_dist': 'Category Distribution',
            'trends': 'Main Trends',
            'analysis': 'Deep Analysis',
            'reco': 'Strategic Recommendations',
            'conclusion': 'Conclusion',
            'obs': 'Observations',
            'vars': 'Variables',
            'num': 'Numeric',
            'cat': 'Categorical',
            'interpretation': 'Interpretation',
        }
    }
    
    t = titles[lang]
    
    # ==================
    # PAGE DE GARDE
    # ==================
    title = doc.add_heading(t['report'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(102, 126, 234)
    title_run.font.size = Pt(28)
    
    # Date
    date_para = doc.add_paragraph(datetime.now().strftime("%d/%m/%Y %H:%M"))
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.runs[0]
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(127, 140, 141)
    
    doc.add_page_break()
    
    # ==================
    # RÉSUMÉ EXÉCUTIF
    # ==================
    doc.add_heading(f"📋 {t['summary']}", 1)
    para = doc.add_paragraph(ai_insights.get('resume_executif', 'N/A'))
    para.paragraph_format.space_after = Pt(12)
    
    # ==================
    # VUE D'ENSEMBLE
    # ==================
    doc.add_heading(f"📊 {t['overview']}", 1)
    
    # Tableau des métriques
    metrics_table = doc.add_table(rows=2, cols=4)
    metrics_table.style = 'Light Grid Accent 1'
    
    # En-têtes
    headers = [t['obs'], t['vars'], t['num'], t['cat']]
    values = [
        str(analysis['shape'][0]),
        str(analysis['shape'][1]),
        str(len(analysis['numeric_cols'])),
        str(len(analysis['categorical_cols']))
    ]
    
    for i, (header, value) in enumerate(zip(headers, values)):
        # En-tête
        cell = metrics_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Valeur
        cell = metrics_table.rows[1].cells[i]
        cell.text = value
        cell.paragraphs[0].runs[0].font.size = Pt(16)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(102, 126, 234)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Espacement
    
    # ==================
    # STATISTIQUES DÉTAILLÉES
    # ==================
    if 'numeric_stats' in analysis and not analysis['numeric_stats'].empty:
        doc.add_page_break()
        doc.add_heading(f"📈 {t['stats']}", 1)
        
        stats_df = analysis['numeric_stats'].T
        
        # Tableau avec 7 colonnes
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        
        # En-têtes
        hdr_cells = table.rows[0].cells
        headers = ['Variable', 'Mean', 'Median', 'Std', 'Min', 'Max', 'Count']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        
        # Données (max 15 variables)
        for col in analysis['numeric_cols'][:15]:
            try:
                stats = analysis['numeric_stats'][col]
                row_cells = table.add_row().cells
                row_cells[0].text = str(col)
                row_cells[1].text = f"{float(stats['mean']):.2f}"
                row_cells[2].text = f"{float(stats['50%']):.2f}"
                row_cells[3].text = f"{float(stats['std']):.2f}"
                row_cells[4].text = f"{float(stats['min']):.2f}"
                row_cells[5].text = f"{float(stats['max']):.2f}"
                row_cells[6].text = f"{int(stats['count'])}"
                
                # Réduire la taille de police
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
            except Exception as e:
                print(f"Error adding stats for {col}: {e}")
    
    # ==================
    # VISUALISATIONS
    # ==================
    if visualizations:
        doc.add_page_break()
        doc.add_heading(f"📊 {t['visualizations']}", 1)
        
        for viz_name, (fig, interpretation) in visualizations.items():
            try:
                # Titre de la visualisation
                viz_title = viz_name.replace('_', ' ').title()
                doc.add_heading(viz_title, 2)
                
                # Ajouter la figure
                add_figure_to_doc(doc, fig, width=6.0)
                
                # Ajouter l'interprétation
                if interpretation:
                    doc.add_heading(f"💡 {t['interpretation']}", 3)
                    interp_para = doc.add_paragraph(interpretation)
                    interp_para.paragraph_format.space_after = Pt(12)
                    
                    # Style italique
                    for run in interp_para.runs:
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(73, 80, 87)
                
                doc.add_paragraph()  # Espacement
                
            except Exception as e:
                print(f"Error adding visualization {viz_name}: {e}")
    
    # ==================
    # DISTRIBUTION DES CATÉGORIES
    # ==================
    if analysis.get('category_dist'):
        doc.add_page_break()
        doc.add_heading(f"🥧 {t['cat_dist']}", 1)
        
        for cat_col, dist in list(analysis['category_dist'].items())[:3]:
            doc.add_heading(cat_col, 2)
            
            # Créer un tableau
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # En-têtes
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Category'
            hdr_cells[1].text = 'Count'
            hdr_cells[2].text = 'Percentage'
            
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Données
            total = sum(dist.values())
            for category, count in list(dist.items())[:10]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(category)
                row_cells[1].text = str(count)
                row_cells[2].text = f"{(count/total*100):.1f}%"
            
            doc.add_paragraph()  # Espacement
    
    # ==================
    # TENDANCES PRINCIPALES
    # ==================
    doc.add_page_break()
    doc.add_heading(f"📈 {t['trends']}", 1)
    
    if ai_insights.get('tendances_principales'):
        for i, trend in enumerate(ai_insights['tendances_principales'], 1):
            para = doc.add_paragraph(f"{i}. {trend}", style='List Number')
            para.paragraph_format.space_after = Pt(6)
    
    # ==================
    # ANALYSES APPROFONDIES
    # ==================
    if ai_insights.get('insights'):
        doc.add_heading(f"🔍 {t['analysis']}", 1)
        
        for insight in ai_insights['insights']:
            doc.add_heading(insight.get('titre', 'N/A'), 2)
            para = doc.add_paragraph(insight.get('description', 'N/A'))
            para.paragraph_format.space_after = Pt(12)
    
    # ==================
    # RECOMMANDATIONS
    # ==================
    doc.add_page_break()
    doc.add_heading(f"💡 {t['reco']}", 1)
    
    if ai_insights.get('recommandations'):
        for i, rec in enumerate(ai_insights['recommandations'], 1):
            # Titre de la recommandation (en gras)
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {rec.get('action', 'N/A')}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(16, 185, 129)
            
            # Justification
            just_para = doc.add_paragraph(rec.get('justification', 'N/A'))
            just_para.paragraph_format.left_indent = Inches(0.3)
            just_para.paragraph_format.space_after = Pt(12)
    
    # ==================
    # CONCLUSION
    # ==================
    doc.add_heading(f"✅ {t['conclusion']}", 1)
    para = doc.add_paragraph(ai_insights.get('conclusion', 'N/A'))
    para.paragraph_format.space_after = Pt(12)
    
    # ==================
    # FOOTER
    # ==================
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    footer = doc.add_paragraph(f"Report generated automatically | Powered by AI")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(108, 117, 125)
    
    date_footer = doc.add_paragraph(datetime.now().strftime("%d %B %Y at %H:%M"))
    date_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_footer.runs[0].font.size = Pt(8)
    date_footer.runs[0].font.color.rgb = RGBColor(149, 165, 166)
    
    # Sauvegarder dans un buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()